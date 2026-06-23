from pathlib import Path
from loguru import logger

try:
    import docx
except ImportError:
    docx = None

from ..base import BaseParser

class DocxParser(BaseParser):
    """Parser for DOCX files using python-docx."""

    def parse(self, file_path: Path, output_dir: Path) -> str:
        if docx is None:
            logger.error("python-docx is not installed.")
            raise ImportError("python-docx is required for DocxParser")

        logger.info(f"Parsing DOCX: {file_path}")
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX {file_path}: {e}")
            raise

        md_lines: list[str] = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # Extract text with inline formatting (bold/italic)
            para_md = ""
            for run in para.runs:
                text = run.text
                if not text.strip():
                    para_md += text
                    continue
                
                # Check formatting
                is_bold = run.bold or (para.style.font.bold if para.style and para.style.font else False)
                is_italic = run.italic or (para.style.font.italic if para.style and para.style.font else False)
                
                # Apply markdown formatting
                # Stripping to avoid formatting spaces: "** text **" is bad, "**text** " is good.
                stripped = text.strip()
                prefix = text[:len(text) - len(text.lstrip())]
                suffix = text[len(text.rstrip()):]
                
                formatted = stripped
                if is_bold and is_italic:
                    formatted = f"***{formatted}***"
                elif is_bold:
                    formatted = f"**{formatted}**"
                elif is_italic:
                    formatted = f"*{formatted}*"
                    
                para_md += f"{prefix}{formatted}{suffix}"
            
            para_md = para_md.strip()
            
            # Apply paragraph-level styles (Headings, Lists)
            style_name = para.style.name.lower() if para.style else ""
            if style_name.startswith('heading'):
                try:
                    level = int(style_name.split()[-1])
                    para_md = f"{'#' * level} {para_md}"
                except ValueError:
                    para_md = f"# {para_md}"
            elif 'List Bullet' in para.style.name: # Original name is case sensitive usually, but docx uses specific names
                para_md = f"- {para_md}"
            elif 'List Number' in para.style.name:
                para_md = f"1. {para_md}"
                
            md_lines.append(para_md)
            
        # Process tables
        for table in doc.tables:
            md_lines.append("\n")
            for i, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells])
                md_lines.append(f"| {row_text} |")
                if i == 0:
                    sep = " | ".join(["---" for _ in row.cells])
                    md_lines.append(f"| {sep} |")
            md_lines.append("\n")

        return "\n\n".join(md_lines)
